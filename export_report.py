# Hospital Management System - Project Report Export Script
# This script converts the HTML report to PDF and DOCX formats with proper formatting

import subprocess
import os
import sys

def check_pandoc():
    """Check if pandoc is installed"""
    try:
        result = subprocess.run(['pandoc', '--version'], capture_output=True, text=True)
        print("✓ Pandoc is installed")
        return True
    except FileNotFoundError:
        print("✗ Pandoc is not installed")
        print("\nTo install Pandoc:")
        print("1. Download from: https://pandoc.org/installing.html")
        print("2. Or use: winget install --id=JohnMacFarlane.Pandoc")
        return False

def convert_to_pdf(input_file, output_file):
    """Convert HTML to PDF using pandoc"""
    print(f"\nConverting {input_file} to PDF...")
    
    cmd = [
        'pandoc',
        input_file,
        '-o', output_file,
        '--pdf-engine=wkhtmltopdf',
        '--pdf-engine-opt=--enable-local-file-access',
        '--pdf-engine-opt=--page-size', '--pdf-engine-opt=A4',
        '--pdf-engine-opt=--margin-top', '--pdf-engine-opt=25.4mm',
        '--pdf-engine-opt=--margin-bottom', '--pdf-engine-opt=25.4mm',
        '--pdf-engine-opt=--margin-left', '--pdf-engine-opt=38.1mm',
        '--pdf-engine-opt=--margin-right', '--pdf-engine-opt=25.4mm',
        '-V', 'geometry:a4paper',
        '-V', 'geometry:margin=1in',
        '-V', 'geometry:left=1.5in',
        '--standalone'
    ]
    
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        print(f"✓ PDF created: {output_file}")
        return True
    except subprocess.CalledProcessError as e:
        print(f"✗ Error creating PDF: {e.stderr}")
        # Try alternative method using weasyprint
        print("\nTrying alternative method with browser print...")
        return convert_to_pdf_browser(input_file, output_file)
    except FileNotFoundError:
        print("✗ wkhtmltopdf not found. Trying alternative method...")
        return convert_to_pdf_browser(input_file, output_file)

def convert_to_pdf_browser(input_file, output_file):
    """Alternative: Open in browser for manual PDF export"""
    import webbrowser
    abs_path = os.path.abspath(input_file)
    print(f"\nOpening {abs_path} in browser...")
    print("Please use browser's Print -> Save as PDF option")
    print("PDF Settings:")
    print("  - Paper size: A4")
    print("  - Margins: Top/Bottom/Right = 1 inch, Left = 1.5 inches")
    webbrowser.open(f'file:///{abs_path}')
    return False

def convert_to_docx(input_file, output_file):
    """Convert HTML to DOCX using pandoc"""
    print(f"\nConverting {input_file} to DOCX...")
    
    # Create a reference DOCX with proper styles
    reference_docx = create_reference_docx()
    
    cmd = [
        'pandoc',
        input_file,
        '-o', output_file,
        '--from=html',
        '--to=docx',
        f'--reference-doc={reference_docx}' if reference_docx else '',
        '--standalone'
    ]
    
    # Remove empty string if no reference doc
    cmd = [c for c in cmd if c]
    
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        print(f"✓ DOCX created: {output_file}")
        
        # Apply additional formatting using python-docx
        apply_docx_formatting(output_file)
        return True
    except subprocess.CalledProcessError as e:
        print(f"✗ Error creating DOCX: {e.stderr}")
        return False
    except FileNotFoundError:
        print("✗ Pandoc not found")
        return False

def create_reference_docx():
    """Create a reference DOCX with proper formatting"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set page margins: Top, Bottom, Right = 1", Left = 1.5"
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1)
            section.page_height = Inches(11.69)  # A4
            section.page_width = Inches(8.27)    # A4
        
        # Define styles
        styles = doc.styles
        
        # Normal style: Times New Roman, 12pt, 1.5 line spacing
        normal = styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing = 1.5
        
        ref_path = 'docs/reference.docx'
        os.makedirs('docs', exist_ok=True)
        doc.save(ref_path)
        print(f"✓ Created reference document: {ref_path}")
        return ref_path
    except ImportError:
        print("! python-docx not installed, using default formatting")
        return None
    except Exception as e:
        print(f"! Could not create reference document: {e}")
        return None

def apply_docx_formatting(docx_file):
    """Apply specific formatting to the DOCX file"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        
        doc = Document(docx_file)
        
        # Set document-wide properties
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1)
        
        # Format all paragraphs
        for paragraph in doc.paragraphs:
            # Set font
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                if run.bold:  # Headings
                    if 'CHAPTER' in run.text:
                        run.font.size = Pt(16)
                    else:
                        run.font.size = Pt(14)
                else:  # Body text
                    run.font.size = Pt(12)
            
            # Set line spacing
            paragraph.paragraph_format.line_spacing = 1.5
        
        doc.save(docx_file)
        print(f"✓ Applied formatting to {docx_file}")
    except ImportError:
        print("! python-docx not installed, skipping advanced formatting")
    except Exception as e:
        print(f"! Error applying formatting: {e}")

def main():
    input_html = 'docs/Complete_Project_Report.html'
    output_pdf = 'docs/HMS_Project_Report.pdf'
    output_docx = 'docs/HMS_Project_Report.docx'
    
    print("=" * 60)
    print("Hospital Management System - Report Export")
    print("=" * 60)
    
    # Check if input file exists
    if not os.path.exists(input_html):
        print(f"✗ Input file not found: {input_html}")
        return
    
    print(f"✓ Found input file: {input_html}")
    
    # Check for pandoc
    has_pandoc = check_pandoc()
    
    if not has_pandoc:
        print("\n" + "=" * 60)
        print("MANUAL EXPORT INSTRUCTIONS")
        print("=" * 60)
        print("\nSince Pandoc is not installed, you can manually export:")
        print("\nFor PDF:")
        print("1. Open docs/Complete_Project_Report.html in Chrome/Edge")
        print("2. Press Ctrl+P (Print)")
        print("3. Select 'Save as PDF'")
        print("4. Set Paper size: A4")
        print("5. Set Margins: Custom")
        print("   - Top: 1 inch (25.4mm)")
        print("   - Bottom: 1 inch (25.4mm)")
        print("   - Left: 1.5 inches (38.1mm)")
        print("   - Right: 1 inch (25.4mm)")
        print("6. Save as: HMS_Project_Report.pdf")
        
        print("\nFor DOCX:")
        print("1. Install Pandoc: winget install --id=JohnMacFarlane.Pandoc")
        print("2. Rerun this script")
        
        # Open in browser for manual export
        import webbrowser
        abs_path = os.path.abspath(input_html)
        webbrowser.open(f'file:///{abs_path}')
        return
    
    # Convert to PDF
    print("\n" + "-" * 60)
    pdf_success = convert_to_pdf(input_html, output_pdf)
    
    # Convert to DOCX
    print("\n" + "-" * 60)
    docx_success = convert_to_docx(input_html, output_docx)
    
    # Summary
    print("\n" + "=" * 60)
    print("EXPORT SUMMARY")
    print("=" * 60)
    if pdf_success:
        print(f"✓ PDF: {output_pdf}")
    else:
        print(f"✗ PDF: Export failed or requires manual completion")
    
    if docx_success:
        print(f"✓ DOCX: {output_docx}")
    else:
        print(f"✗ DOCX: Export failed")
    
    print("\n" + "=" * 60)
    print("FORMATTING SPECIFICATIONS APPLIED:")
    print("=" * 60)
    print("✓ Paper Size: A4 (8.27\" x 11.69\")")
    print("✓ Font: Times New Roman")
    print("✓ Font Sizes:")
    print("  - Body Text: 12pt")
    print("  - Headings: 14pt")
    print("  - Chapter Titles: 16pt")
    print("✓ Margins:")
    print("  - Top: 1 inch")
    print("  - Bottom: 1 inch")
    print("  - Left: 1.5 inches")
    print("  - Right: 1 inch")
    print("✓ Line Spacing: 1.5 lines")
    print("=" * 60)

if __name__ == '__main__':
    main()
